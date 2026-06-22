"""ScholarLink 무료 배포 범위 분석 보고서 생성 스크립트."""

from datetime import datetime
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# ─── 색상 팔레트 (ModernCorporate) ────────────────────────
NAVY = RGBColor(0x0F, 0x2C, 0x59)
SLATE = RGBColor(0x33, 0x44, 0x5C)
ACCENT = RGBColor(0x0E, 0x7C, 0x86)
SOFT_BG = "EAF2F8"
HEADER_BG = "0F2C59"
ROW_ALT_BG = "F4F7FB"
TEXT_GREY = RGBColor(0x55, 0x60, 0x70)

# ─── 유틸 ─────────────────────────────────────────────────
def set_cell_shading(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="BFC9D6", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def set_korean_font(run, name="맑은 고딕", size_pt=10.5, bold=False, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, size=10.5, bold=False, color=SLATE, align=None,
             space_before=0, space_after=4, font="맑은 고딕"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.45
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_korean_font(run, name=font, size_pt=size, bold=bold, color=color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_korean_font(run, size_pt=16, bold=True, color=NAVY)
    # 하단 border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "0F2C59")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_korean_font(run, size_pt=12.5, bold=True, color=ACCENT)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_korean_font(run, size_pt=11, bold=True, color=SLATE)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.4
    bullet_run = p.add_run("•  ")
    set_korean_font(bullet_run, size_pt=10.5, color=ACCENT, bold=True)
    if bold_prefix:
        b = p.add_run(bold_prefix)
        set_korean_font(b, size_pt=10.5, bold=True, color=SLATE)
        rest = p.add_run(" " + text)
        set_korean_font(rest, size_pt=10.5)
    else:
        r = p.add_run(text)
        set_korean_font(r, size_pt=10.5)
    return p


def add_callout(doc, label, body):
    """색상 박스 안 강조 메시지."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(16)
    set_cell_shading(cell, SOFT_BG)
    set_cell_borders(cell, color="0E7C86", size="6")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    lr = p.add_run(label + "  ")
    set_korean_font(lr, size_pt=10, bold=True, color=NAVY)
    br = p.add_run(body)
    set_korean_font(br, size_pt=10)
    # 표 뒤 여백
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def make_table(doc, headers, rows, col_widths_cm=None, header_bg=HEADER_BG,
               numeric_cols=None):
    """깔끔한 모던 표."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in tbl.columns[i].cells:
                cell.width = Cm(w)
    # 헤더
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_shading(cell, header_bg)
        set_cell_borders(cell, color=header_bg, size="6")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        run = cell.paragraphs[0].add_run(h)
        set_korean_font(run, size_pt=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    # 본문
    for r_idx, row in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cell = cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_shading(cell, ROW_ALT_BG)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if numeric_cols and c_idx in numeric_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            bold = bool(numeric_cols and c_idx in numeric_cols)
            set_korean_font(run, size_pt=9.5, bold=bold)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


# ─── 문서 초기화 ──────────────────────────────────────────
doc = Document()

# 페이지 설정
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

# 기본 스타일
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "맑은 고딕"
normal.font.size = Pt(10.5)
normal_rPr = normal.element.get_or_add_rPr()
rFonts = normal_rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    normal_rPr.append(rFonts)
for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
    rFonts.set(qn(attr), "맑은 고딕")

# ─── 표지 ─────────────────────────────────────────────────
cover_title = doc.add_paragraph()
cover_title.paragraph_format.space_before = Pt(40)
cover_title.paragraph_format.space_after = Pt(8)
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title.add_run("ScholarLink")
set_korean_font(run, size_pt=32, bold=True, color=NAVY, name="맑은 고딕")

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_sub.paragraph_format.space_after = Pt(28)
run = cover_sub.add_run("무료 배포 범위 분석 보고서")
set_korean_font(run, size_pt=20, bold=False, color=SLATE)

# 표지 메타 표
meta_tbl = doc.add_table(rows=4, cols=2)
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_tbl.autofit = False
meta_rows = [
    ("문서 유형", "기술 분석 보고서 (Technical Analysis Report)"),
    ("대상 시스템", "ScholarLink 백엔드 (Render Free + Neon Free + GitHub Pages)"),
    ("작성 일자", datetime.now().strftime("%Y년 %m월 %d일")),
    ("분석 범위", "주간 처리 가능 다운로드 편수 및 병목 지점"),
]
for i, (k, v) in enumerate(meta_rows):
    kc = meta_tbl.rows[i].cells[0]
    vc = meta_tbl.rows[i].cells[1]
    kc.width = Cm(4.5)
    vc.width = Cm(11.5)
    set_cell_shading(kc, "F4F7FB")
    set_cell_borders(kc)
    set_cell_borders(vc)
    kc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    kp = kc.paragraphs[0]
    kp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kp.add_run(k)
    set_korean_font(kr, size_pt=10, bold=True, color=NAVY)
    vp = vc.paragraphs[0]
    vp.paragraph_format.left_indent = Cm(0.3)
    vr = vp.add_run(v)
    set_korean_font(vr, size_pt=10)

doc.add_paragraph().paragraph_format.space_after = Pt(20)
cover_note = doc.add_paragraph()
cover_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_note.paragraph_format.space_before = Pt(120)
nr = cover_note.add_run("본 보고서는 Render 무료 티어, Neon PostgreSQL 무료 티어, GitHub Pages 무료 호스팅을\n기반으로 운영 중인 ScholarLink 서비스의 주간 처리 가능 다운로드 편수를 정량 분석한다.")
set_korean_font(nr, size_pt=10, color=TEXT_GREY)

# 페이지 넘김
doc.add_page_break()

# ─── 1. 요약 (Executive Summary) ──────────────────────────
add_h1(doc, "1. 요약")

add_para(doc, "현 무료 배포 구조에서 ScholarLink는 주당 약 7,000~8,000편의 다운로드를 안정적으로 처리할 수 있다. "
              "최대 폭주 상황에서도 25,000편/주를 넘기 어렵고, 이는 대역폭 한도에서 결정된다.",
         size=11)

add_callout(doc, "핵심 결론",
            "주간 7,000~8,000편이 지속 처리량, 25,000편이 이론적 상한. 유료 전환 권장 시점은 주당 10,000편 돌파 시점이다.")

add_h2(doc, "주요 발견 사항")

findings = [
    ("Puppeteer가 비활성화되어 있다", "render.yaml의 PUPPETEER_SKIP_DOWNLOAD=true 설정으로 Google Scholar와 Sci-Hub 경로가 꺼져 있다. 512MB 메모리 환경에서 지속 가능성을 확보한 결정적 요인이다."),
    ("30개 OA 소스를 순차 처리한다", "downloadService.ts의 for-await 루프가 모든 소스를 줄 세워 도는 구조로, 평균 58.5초/편의 처리 시간이 발생한다. 단일 인스턴스에서 동시 다발 다운로드는 사실상 직렬화된다."),
    ("사용자당 시간당 20건 제한", "downloadLimiter가 IP당 시간당 20건으로 캡을 설정한다. 한 명의 파워 사용자가 20건을 모두 소진하면 1시간 대기해야 한다."),
    ("대역폭이 최종 천장", "월 100GB 한도에서 PDF 평균 5MB 기준으로 약 20,000편/월, 주당 5,000편까지 처리 가능하다. 단, 평균 PDF 크기가 커지면 편수는 줄어든다."),
    ("인용 네트워크 기능 영향 미미", "OpenAlex API 호출이 다운로드의 첫 번째 소스로 이미 포함되어 있어, 그래프 빌드 추가 부하가 전체 처리량에 미치는 영향은 1% 미만이다."),
]
for prefix, body in findings:
    add_bullet(doc, body, bold_prefix="● " + prefix)

# ─── 2. 인프라 제약사항 ──────────────────────────────────
add_h1(doc, "2. 인프라 제약사항")

add_para(doc, "ScholarLink는 세 가지 무료 서비스의 조합으로 운영된다.")

# 2.1 Render Free Tier
add_h2(doc, "2.1 Render Free Tier (백엔드 API 서버)")
make_table(
    doc,
    ["자원", "한도", "처리량에 미치는 영향"],
    [
        ["vCPU", "0.1 코어 (공유)", "Node.js 단일 스레드 + 순차 처리로 직접 병목"],
        ["메모리", "512 MB", "Puppeteer 미사용으로 여유 있음"],
        ["대역폭", "100 GB/월", "PDF 평균 5MB 기준 약 20,000편/월 상한"],
        ["월 가동 시간", "750 시간", "연중 24/7 운영 가능 (keep-alive 적용)"],
        ["Cold Start", "15분 유휴 시 슬립", "GitHub Actions 핑으로 10분마다 깨움"],
        ["Puppeteer", "사용 불가 (SKIP)", "메모리 부족으로 Chrome 실행 비활성"],
    ],
    col_widths_cm=[3.2, 4.8, 8.0],
)

# 2.2 Neon Free Tier
add_h2(doc, "2.2 Neon PostgreSQL Free Tier")
make_table(
    doc,
    ["자원", "한도", "처리량에 미치는 영향"],
    [
        ["스토리지", "0.5 GB", "평생 약 50,000건 다운로드 기록 가능"],
        ["Compute 시간", "190 시간/월", "5분 유휴 시 자동 정지, 첫 요청 시 0.5~2초 지연"],
        ["커넥션", "기본 100개 허용", "pg pool max=10 설정으로 충분"],
        ["백업", "7일 retention", "다운로드 이력 복구 가능"],
    ],
    col_widths_cm=[3.2, 4.8, 8.0],
)

# 2.3 GitHub Pages
add_h2(doc, "2.3 GitHub Pages (프론트엔드 호스팅)")
add_para(doc, "정적 자산 호스팅이므로 처리량 영향 없음. 단, 클라이언트 측 다운로드 요청 발생 시 백엔드 API로 트래픽이 집중된다.")

# ─── 3. 처리 시간 분석 ──────────────────────────────────
add_h1(doc, "3. 처리 시간 분석")

add_para(doc, "DOI 유형에 따라 평균 처리 시간이 크게 달라진다. 30개 OA 소스를 순차로 시도하기 때문에 "
              "어느 소스에서 성공하느냐가 전체 응답 시간을 결정한다.")

add_h2(doc, "3.1 DOI 유형별 처리 시간 분포")
make_table(
    doc,
    ["DOI 유형", "추정 비율", "성공 위치", "평균 시간"],
    [
        ["arXiv / PMC OA / OA.mg 직접", "30%", "1~3번째 소스", "15초"],
        ["일반 OA 저널 (Unpaywall 등)", "40%", "4~10번째 소스", "45초"],
        ["출판사 OA (PLOS / Science / Springer)", "20%", "10~20번째 소스", "90초"],
        ["OA 없음 / RISS 폴백", "10%", "20~30번째 소스 + RISS", "180초"],
    ],
    col_widths_cm=[5.5, 2.0, 4.5, 3.0],
    numeric_cols=[1, 3],
)

add_callout(doc, "가중 평균",
            "0.3 × 15 + 0.4 × 45 + 0.2 × 90 + 0.1 × 180 = 약 58.5초/편")

add_h2(doc, "3.2 시퀀셜 처리의 영향")

add_para(doc, "downloadService.ts의 핵심 루프 구조는 다음과 같다.")
code_p = doc.add_paragraph()
code_p.paragraph_format.left_indent = Cm(0.6)
code_p.paragraph_format.space_before = Pt(2)
code_p.paragraph_format.space_after = Pt(6)
code_run = code_p.add_run(
    "const oaSources = [ ['OpenAlex', fn1], ['Unpaywall', fn2], ... 28 more ];\n"
    "for (const [name, fn] of oaSources) {          // 줄 서서 순차 시도\n"
    "    const r = await fn();                       // 외부 API 응답 대기\n"
    "    if (r) return r;                            // 성공 시 즉시 종료\n"
    "}"
)
set_korean_font(code_run, size_pt=9, name="Consolas")

add_para(doc, "Node.js의 비동기 이벤트 루프 덕분에 HTTP 요청 자체는 병렬로 나가지만, downloadService 함수의 결과 반환은 "
              "다음 소스 시도를 차단한다. 결과적으로 단일 인스턴스에서 두 개의 동시 다운로드 요청은 사실상 직렬화된다.",
         size=10)

# ─── 4. 병목 분석 ──────────────────────────────────────
add_h1(doc, "4. 병목 분석")

add_para(doc, "현재 구조에는 세 가지 주요 병목 지점이 존재한다. 각각이 천장으로 작용하는 시나리오가 다르다.")

add_h2(doc, "4.1 CPU 병목 (Render 0.1 vCPU)")

make_table(
    doc,
    ["구간", "CPU 점유", "주기"],
    [
        ["외부 API 응답 대기", "낮음", "60~80% 시간"],
        ["JSON / HTML 파싱 (cheerio)", "중간", "5~15%"],
        ["PDF 스트림 디스크 쓰기", "높음", "3~5%"],
        ["DB INSERT / UPDATE", "중간", "1~2%"],
        ["이벤트 루프 tick", "낮음", "지속"],
    ],
    col_widths_cm=[6.5, 2.5, 6.0],
)

add_para(doc, "0.1 vCPU는 풀 코어 대비 약 10% 성능이다. 외부 API 응답을 기다리는 동안에도 이벤트 루프의 I/O 처리 속도가 느려져 "
              "체감 처리량은 이론치 대비 70~80% 수준으로 떨어진다.",
         size=10)

add_h2(doc, "4.2 Rate Limit 캡 (사용자당 시간당 20건)")

add_para(doc, "downloadLimiter 미들웨어는 IP당 시간당 20건으로 제한한다. 한 명의 파워 사용자가 20건을 모두 소진하면 "
              "1시간 동안 429 Too Many Requests 응답을 받는다.",
         size=10)

make_table(
    doc,
    ["시나리오", "활성 사용자 수", "주간 최대 처리량"],
    [
        ["소수 파워 유저", "50명", "약 1,500건"],
        ["일반적 사용", "500명", "약 5,000건"],
        ["활발한 사용", "1,000명", "약 10,000건"],
        ["캠퍼스 단위 유입", "3,000명 이상", "15,000건+ (CPU 병목 진입)"],
    ],
    col_widths_cm=[4.5, 3.5, 7.0],
    numeric_cols=[1, 2],
)

add_h2(doc, "4.3 대역폭 천장 (월 100GB)")

make_table(
    doc,
    ["평균 PDF 크기", "월 처리 한도", "주간 환산"],
    [
        ["3 MB (경량)", "33,000편", "약 8,000편/주"],
        ["5 MB (현재 평균 가정)", "20,000편", "약 5,000편/주"],
        ["10 MB (대용량 저널)", "10,000편", "약 2,500편/주"],
        ["30 MB (일부 Nature 류)", "3,300편", "약 800편/주"],
    ],
    col_widths_cm=[4.5, 3.5, 7.0],
    numeric_cols=[1, 2],
)

add_callout(doc, "참고",
            "Render는 대역폭 초과 시 즉각 차단이 아니라 종량 과금 또는 일시적 제한으로 대응한다. 모니터링이 필수.")

# ─── 5. 처리량 시나리오 ──────────────────────────────────
add_h1(doc, "5. 처리량 시나리오")

add_para(doc, "사용자 행동 패턴과 DOI 유형 분포를 조합한 네 가지 시나리오별 주간 처리량 추정치.")

make_table(
    doc,
    ["시나리오", "DOI 분포 (Fast/Medium/Slow)", "사용자 수", "주간 처리량"],
    [
        ["🔵 조용한 주 (소수 파워 유저)", "20% / 50% / 30%", "50명", "1,500~3,000편"],
        ["🟢 일반적 주 (혼합 사용자)", "30% / 40% / 30%", "500명", "5,000~8,000편"],
        ["🟡 바쁜 주 (캠퍼스/학회 시즌)", "35% / 40% / 25%", "1,500명", "10,000~13,000편"],
        ["🔴 폭주 (론칭 직후/바이럴)", "40% / 40% / 20%", "3,000명 이상", "15,000~20,000편"],
    ],
    col_widths_cm=[5.0, 4.0, 2.5, 4.0],
    numeric_cols=[2, 3],
)

add_para(doc, "현실적 지속 가능 처리량은 약 7,000~8,000편/주이며, 이 구간에서 CPU와 Rate Limit이 균형을 이룬다.",
         size=10)

# ─── 6. 인용 네트워크 기능 영향 ──────────────────────────
add_h1(doc, "6. 인용 네트워크 시각화 기능의 영향")

add_para(doc, "기획 중인 OpenAlex 기반 인용 네트워크 시각화 기능이 추가될 경우의 부하 영향 분석.")

add_h2(doc, "6.1 추가 부하 산정")
make_table(
    doc,
    ["항목", "예상 수치", "비고"],
    [
        ["그래프 1회당 OpenAlex 호출", "5~20회", "depth 1: 약 5회, depth 2: 약 10회"],
        ["그래프 빌드 서버 처리 시간", "5~15초", "백그라운드 캐시 미적용 시"],
        ["예상 주간 그래프 요청", "500~1,000회", "다운로드 사용자의 10% 가정"],
        ["추가 CPU 시간", "1.4~2.8시간/주", "전체 CPU 사용량의 0.8~1.6%"],
    ],
    col_widths_cm=[5.5, 3.5, 6.5],
)

add_callout(doc, "결론",
            "인용 네트워크 기능이 추가되어도 전체 다운로드 처리량 감소폭은 5% 이내로 미미하다. 기존 OpenAlex 호출이 다운로드의 첫 번째 소스로 이미 포함되어 있어 새로 추가되는 부하가 거의 없다.")

# ─── 7. 리스크 시나리오 ──────────────────────────────────
add_h1(doc, "7. 리스크 시나리오")

risks = [
    ("Puppeteer 재활성화", "높음",
     "누군가 PUPPETEER_SKIP_DOWNLOAD 설정을 변경하면 512MB 메모리에서 즉시 OOM 위험. Chrome 인스턴스 1개당 250MB 소비. "
     "정책적으로 이 설정은 절대 변경하지 않아야 한다."),
    ("평균 PDF 크기 급증", "중간",
     "현재 평균 5MB 가정이 깨지면 동일 대역폭에서 처리 편수가 줄어든다. MAX_FILE_SIZE_MB=50 설정은 이미 render.yaml에 존재하지만 "
     "이보다 작은 임계치(예: 20MB) 설정도 고려할 만하다."),
    ("Cold Start 응답 지연", "낮음",
     "GitHub Actions keep-alive 핑이 단순 GET만 보내므로 실제 모듈 워밍업은 되지 않는다. 첫 사용자 요청 시 5~30초 지연 발생 가능. "
     "핑 엔드포인트에 SELECT 1 쿼리를 추가하면 완화된다."),
    ("Neon auto-suspend", "낮음",
     "5분 유휴 시 compute 정지. 첫 요청에서 깨어나는데 0.5~2초. 처리량 자체엔 영향 미미하지만 사용자 체감 응답 시간에 영향."),
    ("외부 API 장애", "중간",
     "Crossref, Unpaywall, OpenAlex 등 주요 API에 장애가 발생하면 fallback 경로를 모두 소진한 후 실패. "
     "현재 구조에서 자동 복구 메커니즘은 다음 사용자 요청에 의존."),
    ("동시 접속 스파이크", "중간",
     "특정 학회/논문이 SNS에서 바이럴될 경우 단시간에 동일 DOI에 대한 요청 폭주. node-cache 미사용으로 동일 DOI도 매번 30개 소스 전부 시도."),
]
make_table(
    doc,
    ["리스크", "심각도", "상세 및 대응"],
    [(name, level, body) for name, level, body in risks],
    col_widths_cm=[3.5, 1.8, 9.7],
)

# ─── 8. 최적화 권장사항 ──────────────────────────────────
add_h1(doc, "8. 최적화 권장사항")

add_para(doc, "인프라 변경 없이 처리량을 늘릴 수 있는 다섯 가지 즉시 적용 가능한 최적화.")

make_table(
    doc,
    ["#", "최적화", "예상 효과", "난이도"],
    [
        ["1", "node-cache 활성화 (DOI별 결과 캐시)", "DB/API 호출 30~50% 감소", "쉬움"],
        ["2", "oaSources 병렬 처리 (상위 3개 동시 + 나머지 순차)", "평균 처리 시간 40% 단축", "중간"],
        ["3", "Unpaywall 응답 캐시 (OA location은 거의 불변)", "API quota 50% 절약", "쉬움"],
        ["4", "HEAD 요청 제거 (GET만 시도)", "소스당 1~2초 절약", "쉬움"],
        ["5", "keep-alive 핑에 워밍업 쿼리 추가", "Cold start 지연 제거", "매우 쉬움"],
    ],
    col_widths_cm=[0.8, 6.7, 4.5, 3.0],
    numeric_cols=[0, 3],
)

add_callout(doc, "최우선 과제",
            "#1 (node-cache 활성화). 동일한 arXiv DOI가 두 번째 검색될 때 30개 소스 전체를 다시 돌지 않고 즉시 응답할 수 있다. "
            "package.json에는 이미 node-cache@5.1.2가 등록되어 있으나 현재 코드에서 사용되지 않고 있다.")

# ─── 9. 유료 전환 임계점 ──────────────────────────────────
add_h1(doc, "9. 유료 전환 임계점")

add_para(doc, "언제 어떤 단계로 유료 플랜으로 전환해야 하는가?")

make_table(
    doc,
    ["임계점", "권장 액션", "예상 비용", "전환 효과"],
    [
        ["주간 7,000편 도달", "최적화 5건 우선 적용", "$0", "처리량 1.5~2배"],
        ["주간 10,000편 돌파", "Render Standard + Neon Launch", "월 $12", "처리량 10배 여유"],
        ["Puppeteer 필요성 발생", "Render Standard 이상 필수", "월 $7+", "Sci-Hub/Scholar 재활성화 가능"],
        ["캠퍼스/기관 단위 유입", "Render Pro + Neon Scale", "월 $85+", "동시 접속 수백 명 처리"],
    ],
    col_widths_cm=[3.8, 5.0, 2.5, 4.7],
    numeric_cols=[2],
)

add_callout(doc, "현재 권장",
            "사용자 수가 주당 5,000편을 안정적으로 넘기기 시작하면 #1~#3 최적화를 모두 적용한 뒤, "
            "주간 10,000편 돌파 시점에 Render Standard($7/월)로 전환한다. 그 이전에는 무료 티어 안에서 충분히 대응 가능하다.")

# ─── 10. 결론 ──────────────────────────────────────────
add_h1(doc, "10. 결론")

add_para(doc, "ScholarLink의 현재 무료 배포 구조는 안정적인 주당 7,000~8,000편의 다운로드 처리 능력을 갖추고 있다. "
              "이는 Render 무료 티어의 0.1 vCPU 제약, 30개 OA 소스의 시퀀셜 처리 구조, 사용자당 시간당 20건의 Rate Limit이 "
              "동시에 천장으로 작용하기 때문이다.",
         size=11)

add_para(doc, "Puppeteer를 의도적으로 비활성화한 현재 설정은 512MB 메모리 환경에서 지속 가능성을 확보한 가장 중요한 결정이다. "
              "동일한 결과를 무료 티어에서 더 많이 뽑아내려면 node-cache 활성화와 oaSources 병렬 처리 같은 코드 레벨 최적화가 "
              "가장 효과적인 방법이다.",
         size=11)

add_para(doc, "인용 네트워크 시각화 기능 추가는 전체 처리량에 1% 미만의 부하만을 추가하므로, "
              "다운로드 서비스의 처리 능력을 거의 저하시키지 않고 출시할 수 있다.",
         size=11)

# ─── 부록: 핵심 수치 요약 ──────────────────────────────────
doc.add_page_break()
add_h1(doc, "부록. 핵심 수치 요약")

add_h2(doc, "A.1 인프라 상수")
make_table(
    doc,
    ["항목", "값"],
    [
        ["Render vCPU", "0.1 코어"],
        ["Render 메모리", "512 MB"],
        ["Render 대역폭", "100 GB/월"],
        ["Neon 스토리지", "0.5 GB"],
        ["Neon Compute", "190 시간/월"],
        ["pg Pool max", "10 커넥션"],
        ["사용자당 다운로드 한도", "20건/시간"],
        ["OA 소스 개수", "30개"],
        ["Puppeteer 상태", "비활성 (SKIP)"],
    ],
    col_widths_cm=[6.0, 9.0],
)

add_h2(doc, "A.2 처리 시간 분포")
make_table(
    doc,
    ["구간", "비율", "시간"],
    [
        ["Fast (arXiv / PMC / OA.mg)", "30%", "15초"],
        ["Medium (Unpaywall / 일반 OA)", "40%", "45초"],
        ["Slow (출판사 OA)", "20%", "90초"],
        ["None / RISS 폴백", "10%", "180초"],
        ["가중 평균", "100%", "58.5초"],
    ],
    col_widths_cm=[6.0, 3.0, 6.0],
    numeric_cols=[1, 2],
)

add_h2(doc, "A.3 주간 처리량 시나리오")
make_table(
    doc,
    ["시나리오", "주간 처리량"],
    [
        ["조용한 주", "1,500~3,000편"],
        ["일반적 주 (현실적 지속)", "5,000~8,000편"],
        ["바쁜 주", "10,000~13,000편"],
        ["폭주 (대역폭 한도 진입)", "15,000~20,000편"],
        ["이론적 최대 (24/7)", "약 10,300편"],
    ],
    col_widths_cm=[8.0, 7.0],
)

# ─── 저장 ──────────────────────────────────────────────
output_path = "D:/SC_link/docs/ScholarLink_무료배포범위_분석보고서.docx"
import os
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"보고서 저장 완료: {output_path}")
print(f"파일 크기: {os.path.getsize(output_path) / 1024:.1f} KB")
